import string
from secrets import choice
from docx import Document
import flet as ft

# Генерация паролей
def password_gen(length):
    alphabet = string.ascii_letters + string.digits + string.punctuation
    return ''.join(choice(alphabet) for _ in range(length))

# Сохранение в .docx файл
def save_to_docx(passwords, filename='passwords.docx'):
    doc = Document()
    doc.add_heading('Generated Passwords', level=1)
    for password in passwords:
        doc.add_paragraph(password)
    doc.save(filename)

# Основной GUI
def main(page: ft.Page):
    page.title = "Cyberpunk Password Generator"
    page.window.width = 600
    page.window.height = 800
    page.theme_mode = ft.ThemeMode.DARK

    # Цвета и стили
    bg_color = ft.Colors.BLACK
    text_color = ft.Colors.YELLOW_400
    accent_color = ft.Colors.CYAN_400

    page.bgcolor = bg_color
    page.scroll = ft.ScrollMode.AUTO

    # Список паролей
    passwords = []

    # Обновление паролей в таблице
    def update_table():
        table.controls.clear()
        for idx, password in enumerate(passwords):
            table.controls.append(ft.Text(f"{idx + 1}. {password}", color=text_color))

    # Генерация паролей
    def generate_passwords(e):
        nonlocal passwords
        if not length_field.value.isdigit() or int(length_field.value) not in [10, 12, 15]:
            page.snack_bar = ft.SnackBar(ft.Text("Invalid length! Choose 10, 12, or 15.", color=accent_color))
            page.snack_bar.open = True
            page.update()
            return

        length = int(length_field.value)
        passwords = [password_gen(length) for _ in range(10)]
        update_table()
        page.update()

    # Сохранение паролей
    def save_passwords(e):
        if passwords:
            save_to_docx(passwords)
            page.snack_bar = ft.SnackBar(ft.Text("Passwords saved to 'passwords.docx'.", color=accent_color))
        else:
            page.snack_bar = ft.SnackBar(ft.Text("No passwords to save!", color=accent_color))
        page.snack_bar.open = True
        page.update()

    # Поле для ввода длины
    length_field = ft.TextField(label="Password Length (10, 12, 15)", color=text_color, bgcolor=bg_color)

    # Кнопки
    generate_button = ft.ElevatedButton("Generate Passwords", on_click=generate_passwords, bgcolor=accent_color)
    save_button = ft.ElevatedButton("Save to File", on_click=save_passwords, bgcolor=accent_color)

    # Таблица сгенерированных паролей
    table = ft.Column(scroll=ft.ScrollMode.AUTO, expand=True)

    # Основной макет
    page.add(
        ft.Column(
            [
                ft.Text("Password Generator", size=24, color=text_color, weight=ft.FontWeight.BOLD),
                length_field,
                ft.Row([generate_button, save_button], alignment=ft.MainAxisAlignment.SPACE_EVENLY),
                ft.Divider(),
                ft.Text("Generated Passwords:", color=text_color, size=20, weight=ft.FontWeight.BOLD),
                table,
            ],
            spacing=20,
            alignment=ft.MainAxisAlignment.START,
            expand=True,
        )
    )

# Запуск приложения
if __name__ == "__main__":
    ft.app(target=main)